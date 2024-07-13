import os
import re
import string
import subprocess

import openai
import json
import requests
import time
import random
from PIL import Image
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from selenium import webdriver
from docx.shared import Inches
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from selenium.common.exceptions import WebDriverException
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager

chrome_driver_path = './chromedriver.exe'
service = Service(executable_path=chrome_driver_path)

option = webdriver.ChromeOptions()
option.add_experimental_option("excludeSwitches", ['enable-automation'])
option.add_experimental_option("detach", True)
option.add_argument('--disable-blink-features=AutomationControlled')  # 伪装浏览器不被selenium检测
option.add_argument('--disable-gpu')  # 对于某些系统，这可能有帮助
option.add_argument('--no-sandbox')  # 绕过操作系统安全模型

# option.add_argument('--incognito')# 添加无痕模式参数
# option.add_argument('--headless') #不显示浏览器
# option.add_argument('--disable-gpu')

"""

这里写产品描述



"""

urlHrefARR = []  # 生成的文章链接
titleArr = []  # 生成的文章标题
imgType = None
user_profit = []


# imgType=''

def get_chatgpt(content):
    openai.api_key = "sk-7Dsj4P3kjdS4h20z40D8F4B5CfD640Ec8b231aA97aC32cAd"
    openai.base_url = "https://free.gpt.ge/v1/"
    openai.default_headers = {"x-foo": "true"}
    completion = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "user",
                "content": content
            },
        ],
    )

    print(' = ' * 5 + " 文章处理完成%s " % (' = ' * 5))
    return completion.choices[0].message.content


# 文心一言

def get_content_baidu(content):
    # 修改成自己的api key和secret key
    url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/eb-instant?access_token=" + get_access_token()
    # 注意message必须是奇数条
    payload = json.dumps({
        "messages": [
            {
                "role": "user",
                "content": content
            }
        ]
    })
    headers = {
        'Content-Type': 'application/json'
    }
    res = requests.request("POST", url, headers=headers, data=payload).json()
    return res['result']


def get_access_token():
    API_KEY = "rDJ6r8SHWR9KXX8Qs1AVLI2s"
    SECRET_KEY = "Y9jUgM7gX2wwIp4gVnhWNOI9WMtk3qIM"
    url = "https://aip.baidubce.com/oauth/2.0/token"
    params = {"grant_type": "client_credentials", "client_id": API_KEY, "client_secret": SECRET_KEY}
    return str(requests.post(url, params=params).json().get("access_token"))


def set_sessid(cateArr):
    response = requests.get('https://www.baoww.net/index', )
    server_cookies = response.cookies
    server_cookies_dict = requests.utils.dict_from_cookiejar(server_cookies)
    time.sleep(1)
    cookies = {
        'PHPSESSID': str(server_cookies_dict['PHPSESSID']),
    }
    headers = {
        'authority': 'www.baoww.net',
        'accept': '*/*',
        'accept-language': 'zh-CN,zh;q=0.9,en;q=0.8,en-GB;q=0.7,en-US;q=0.6',
        'cache-control': 'no-cache',
        'content-type': 'application/x-www-form-urlencoded; charset=UTF-8',
        # 'cookie': 'PHPSESSID=fccpl72uj59cdtfp6tf58tj4s3; plat=bw; ver=ok; Hm_lvt_62f4222b024b771cac70fdaee5c8bb75=1720850685; Hm_lpvt_62f4222b024b771cac70fdaee5c8bb75=1720850685; HMACCOUNT=DFB522E4F78DA7E9',
        'origin': 'https://www.baoww.net',
        'pragma': 'no-cache',
        'referer': 'https://www.baoww.net/login',
        'sec-ch-ua': '"Chromium";v="122", "Not(A:Brand";v="24", "Microsoft Edge";v="122"',
        'sec-ch-ua-mobile': '?0',
        'sec-ch-ua-platform': '"Windows"',
        'sec-fetch-dest': 'empty',
        'sec-fetch-mode': 'cors',
        'sec-fetch-site': 'same-origin',
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36 Edg/122.0.0.0',
        'x-requested-with': 'XMLHttpRequest',
    }

    data = {
        'entryname': '15550951182',
        'pass': 'qwe123123',
    }
    response = requests.post('https://www.baoww.net/dologin', cookies=cookies, headers=headers, data=data)
    if response.json()['success']:
        print('登录成功')
        get_content_url(cateArr, server_cookies_dict['PHPSESSID'])
    else:
        print('登录失败')
        return None


# 获取所有热点文章和链接
def get_content_url(cateArr, get_phpsessid):
    headers = {
        'Referer': 'https://www.baoww.net/index',
        'Upgrade-Insecure-Requests': '1',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36',
        'sec-ch-ua': '"Not/A)Brand";v="8", "Chromium";v="126", "Google Chrome";v="126"',
        'sec-ch-ua-mobile': '?0',
        'sec-ch-ua-platform': '"Windows"',
    }
    cookies = {
        'PHPSESSID': str(get_phpsessid),
    }
    platLiat = ['tou']
    params = {
        'plat': platLiat,
        'cate': cateArr,
        # 'sort': 'reads'
    }
    response = requests.get('https://www.baoww.net/index', params=params, cookies=cookies, headers=headers)
    res = BeautifulSoup(response.text, 'html.parser')
    urlList = res.find_all('div', class_='middle aligned content')
    count = 0
    for url in urlList:
        # if count >= random.randint(3, 5):  # 修改这个值以获取更多或更少的 URL
        #     break
        urlHref = url.find('a')['href']
        urlHrefARR.append(urlHref)
        count += 1
    time.sleep(random.randint(2, 6))
    random.shuffle(urlHrefARR)
    if urlHrefARR:
        pass
    else:
        print("未获取到可用的文章已自动停止该项目的执行")
        time.sleep(4)
        exit(1)


# 保存文章内容
def goto_content(url):
    cookies = {
        '__ac_signature': '_02B4Z6wo00f01li2QHAAAIDDdjGNW13cF85YkkTAAPBqxxZPGJSh8S2uYbp142zu0slNQjGRP.C9jVsdADv4x3btLmWZKobiYmx.uXEpcS-5qy74f30.vROVK6J3a39-NhX1ShxGi92uEUYF16',
        'tt_webid': '7383222509079168524',
        's_v_web_id': 'verify_lxpsbcu1_tqLAVkC4_0dBJ_46FJ_9nPi_9HBDhEAC4VUS',
        'local_city_cache': '%E5%8C%97%E4%BA%AC',
        'csrftoken': 'b9d561d078967fd183f009747d1de3c1',
        '_ga': 'GA1.1.493049990.1719040480',
        'passport_csrf_token': '22cb2225bd54604b1c5111feb744baf1',
        'passport_csrf_token_default': '22cb2225bd54604b1c5111feb744baf1',
        'n_mh': '_m2RIZ3NCWiEvhGvMJWu7G75-qDajYEQe1a7Z_YR7EM',
        'sso_uid_tt': '8896d87d65bb7936cb16c48a8320f260',
        'sso_uid_tt_ss': '8896d87d65bb7936cb16c48a8320f260',
        'toutiao_sso_user': '43eea186639cc12bbcbd32bb3903050a',
        'toutiao_sso_user_ss': '43eea186639cc12bbcbd32bb3903050a',
        'sid_ucp_sso_v1': '1.0.0-KDEyYmU3YjM5NzgxMzRkNTM5NzY0M2FmZDY3YTMzNThhNTdjNzUzMzkKHgj4qJDRl8zGARDU8uSzBhgYIAwwyMXeqAY4BkD0BxoCaGwiIDQzZWVhMTg2NjM5Y2MxMmJiY2JkMzJiYjM5MDMwNTBh',
        'ssid_ucp_sso_v1': '1.0.0-KDEyYmU3YjM5NzgxMzRkNTM5NzY0M2FmZDY3YTMzNThhNTdjNzUzMzkKHgj4qJDRl8zGARDU8uSzBhgYIAwwyMXeqAY4BkD0BxoCaGwiIDQzZWVhMTg2NjM5Y2MxMmJiY2JkMzJiYjM5MDMwNTBh',
        'passport_auth_status': '3717946524980dca00355f9bba82f6ed%2C',
        'passport_auth_status_ss': '3717946524980dca00355f9bba82f6ed%2C',
        'sid_guard': '5ddb877b1cbd222e24d6687f62b4a245%7C1719220566%7C5184000%7CFri%2C+23-Aug-2024+09%3A16%3A06+GMT',
        'uid_tt': '7d380573636fff9edfcbe8ee19017d0d',
        'uid_tt_ss': '7d380573636fff9edfcbe8ee19017d0d',
        'sid_tt': '5ddb877b1cbd222e24d6687f62b4a245',
        'sessionid': '5ddb877b1cbd222e24d6687f62b4a245',
        'sessionid_ss': '5ddb877b1cbd222e24d6687f62b4a245',
        'sid_ucp_v1': '1.0.0-KDE3NGY0NTk4MGY1ZTg2ZGI5MDBjYWUwYjAwYzAzNmFhM2M1NGZjNzAKGAj4qJDRl8zGARDW8uSzBhgYIAw4BkD0BxoCbHEiIDVkZGI4NzdiMWNiZDIyMmUyNGQ2Njg3ZjYyYjRhMjQ1',
        'ssid_ucp_v1': '1.0.0-KDE3NGY0NTk4MGY1ZTg2ZGI5MDBjYWUwYjAwYzAzNmFhM2M1NGZjNzAKGAj4qJDRl8zGARDW8uSzBhgYIAw4BkD0BxoCbHEiIDVkZGI4NzdiMWNiZDIyMmUyNGQ2Njg3ZjYyYjRhMjQ1',
        'store-region': 'sg',
        'store-region-src': 'uid',
        'odin_tt': '95ba4b9e10c65d6f0c244c3b646586700d0de86a5ed3b5f4c6437683f478a2fca912cb92a186c28e8b155e2353dff560',
        'gfkadpd': '24,6457',
        '_ga_QEHZPBE5HH': 'GS1.1.1719389456.6.1.1719389762.0.0.0',
        'tt_anti_token': 'sqOXwxlSA6b-14643258799722d00cdd4893a777b86e0f6540852a4c8ef4b30a1d28af54aee7',
        'tt_scid': 'cLbbBvOqoGIrxjTpcEhhxfXdyb5RXOF3mIrgdQw7r2ly19M9Am8wSnF2YMoht9qq605e',
        'ttwid': '1%7CHk-KA8P-lOGkaVkHmttlCZVnUnAD-GoZoYShHFta2KI%7C1719389760%7C2100147d69dd94978c04286550d09d7f266d443db69793858d33fbda4a3fd093',
        'msToken': 'xPSazmywG7kbq9K6tRMgV2ssRfyuyeKW0hWsAqJzIQZwI6VH2sbHme5956SG6n9SVSSS81R74KjDco3HVu4-_ZOsfw5_XyDfgUtUTFDbU1AzTfaqMWlT',
    }
    headers = {
        'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7',
        'accept-language': 'zh-CN,zh;q=0.9',
        'cache-control': 'max-age=0',
        # 'cookie': '__ac_signature=_02B4Z6wo00f01li2QHAAAIDDdjGNW13cF85YkkTAAPBqxxZPGJSh8S2uYbp142zu0slNQjGRP.C9jVsdADv4x3btLmWZKobiYmx.uXEpcS-5qy74f30.vROVK6J3a39-NhX1ShxGi92uEUYF16; tt_webid=7383222509079168524; s_v_web_id=verify_lxpsbcu1_tqLAVkC4_0dBJ_46FJ_9nPi_9HBDhEAC4VUS; local_city_cache=%E5%8C%97%E4%BA%AC; csrftoken=b9d561d078967fd183f009747d1de3c1; _ga=GA1.1.493049990.1719040480; passport_csrf_token=22cb2225bd54604b1c5111feb744baf1; passport_csrf_token_default=22cb2225bd54604b1c5111feb744baf1; n_mh=_m2RIZ3NCWiEvhGvMJWu7G75-qDajYEQe1a7Z_YR7EM; sso_uid_tt=8896d87d65bb7936cb16c48a8320f260; sso_uid_tt_ss=8896d87d65bb7936cb16c48a8320f260; toutiao_sso_user=43eea186639cc12bbcbd32bb3903050a; toutiao_sso_user_ss=43eea186639cc12bbcbd32bb3903050a; sid_ucp_sso_v1=1.0.0-KDEyYmU3YjM5NzgxMzRkNTM5NzY0M2FmZDY3YTMzNThhNTdjNzUzMzkKHgj4qJDRl8zGARDU8uSzBhgYIAwwyMXeqAY4BkD0BxoCaGwiIDQzZWVhMTg2NjM5Y2MxMmJiY2JkMzJiYjM5MDMwNTBh; ssid_ucp_sso_v1=1.0.0-KDEyYmU3YjM5NzgxMzRkNTM5NzY0M2FmZDY3YTMzNThhNTdjNzUzMzkKHgj4qJDRl8zGARDU8uSzBhgYIAwwyMXeqAY4BkD0BxoCaGwiIDQzZWVhMTg2NjM5Y2MxMmJiY2JkMzJiYjM5MDMwNTBh; passport_auth_status=3717946524980dca00355f9bba82f6ed%2C; passport_auth_status_ss=3717946524980dca00355f9bba82f6ed%2C; sid_guard=5ddb877b1cbd222e24d6687f62b4a245%7C1719220566%7C5184000%7CFri%2C+23-Aug-2024+09%3A16%3A06+GMT; uid_tt=7d380573636fff9edfcbe8ee19017d0d; uid_tt_ss=7d380573636fff9edfcbe8ee19017d0d; sid_tt=5ddb877b1cbd222e24d6687f62b4a245; sessionid=5ddb877b1cbd222e24d6687f62b4a245; sessionid_ss=5ddb877b1cbd222e24d6687f62b4a245; sid_ucp_v1=1.0.0-KDE3NGY0NTk4MGY1ZTg2ZGI5MDBjYWUwYjAwYzAzNmFhM2M1NGZjNzAKGAj4qJDRl8zGARDW8uSzBhgYIAw4BkD0BxoCbHEiIDVkZGI4NzdiMWNiZDIyMmUyNGQ2Njg3ZjYyYjRhMjQ1; ssid_ucp_v1=1.0.0-KDE3NGY0NTk4MGY1ZTg2ZGI5MDBjYWUwYjAwYzAzNmFhM2M1NGZjNzAKGAj4qJDRl8zGARDW8uSzBhgYIAw4BkD0BxoCbHEiIDVkZGI4NzdiMWNiZDIyMmUyNGQ2Njg3ZjYyYjRhMjQ1; store-region=sg; store-region-src=uid; odin_tt=95ba4b9e10c65d6f0c244c3b646586700d0de86a5ed3b5f4c6437683f478a2fca912cb92a186c28e8b155e2353dff560; gfkadpd=24,6457; _ga_QEHZPBE5HH=GS1.1.1719389456.6.1.1719389762.0.0.0; tt_anti_token=sqOXwxlSA6b-14643258799722d00cdd4893a777b86e0f6540852a4c8ef4b30a1d28af54aee7; tt_scid=cLbbBvOqoGIrxjTpcEhhxfXdyb5RXOF3mIrgdQw7r2ly19M9Am8wSnF2YMoht9qq605e; ttwid=1%7CHk-KA8P-lOGkaVkHmttlCZVnUnAD-GoZoYShHFta2KI%7C1719389760%7C2100147d69dd94978c04286550d09d7f266d443db69793858d33fbda4a3fd093; msToken=xPSazmywG7kbq9K6tRMgV2ssRfyuyeKW0hWsAqJzIQZwI6VH2sbHme5956SG6n9SVSSS81R74KjDco3HVu4-_ZOsfw5_XyDfgUtUTFDbU1AzTfaqMWlT',
        'priority': 'u=0, i',
        'sec-ch-ua': '"Not/A)Brand";v="8", "Chromium";v="126", "Google Chrome";v="126"',
        'sec-ch-ua-mobile': '?0',
        'sec-ch-ua-platform': '"Windows"',
        'sec-fetch-dest': 'document',
        'sec-fetch-mode': 'navigate',
        'sec-fetch-site': 'none',
        'sec-fetch-user': '?1',
        'upgrade-insecure-requests': '1',
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36',
    }
    res = requests.get(url, cookies=cookies, headers=headers)
    return res.text


# 裁剪图片
def crop_bottom_pixels(image_path, pixels_to_crop, output_path):
    """
    裁剪图片底部指定像素数的区域。

    :param image_path: 图片的路径
    :param pixels_to_crop: 从底部裁剪的像素数
    :param output_path: 裁剪后图片的保存路径
    """
    # 打开图片
    img = Image.open(image_path)

    # 获取图片的宽度和高度
    width, height = img.size

    # 计算裁剪区域：保留图片的全部宽度，从顶部开始到底部留出要裁剪的像素数
    crop_area = (0, 0, width, height - pixels_to_crop)

    # 根据给定的裁剪区域裁剪图片
    cropped_img = img.crop(crop_area)

    # 保存裁剪后的图片
    cropped_img.save(output_path)


# 保存图片
def download_image(url, i, ):
    global imgType
    """
    下载URL中的图片并保存到指定路径。

    :param url: 图片的URL地址
    :param save_path: 本地保存图片的路径（包括文件名和扩展名）
    """
    try:
        # 发送HTTP GET请求获取图片数据
        response = requests.get(url, stream=True)
        content_type = response.headers.get("Content-Type")
        # 检查请求是否成功
        response.raise_for_status()
        imgType = str(re.search(r'image\/(.*)', content_type).group(1))
        if imgType == 'gif':
            # 保存图片到指定路径
            image = BytesIO(response.content)
            with Image.open(image) as im:
                im.save('./img/' + str(i) + '.' + imgType, save_all=True)
        else:
            if imgType == 'webp':
                imgType = 'jpg'
            # 使用BytesIO作为临时存储，避免直接写入文件系统
            image = Image.open(BytesIO(response.content))
            # 保存图片到指定路径
            image.save('./img/' + str(i) + '.' + imgType)
            time.sleep(1)
            crop_bottom_pixels('./img/' + str(i) + '.' + imgType, 35, './img/' + str(i) + '.' + imgType)
        return imgType
    except requests.RequestException as e:
        print(f"下载图片时发生错误：{e}")


# 处理文章和图片
def handle_content_img():
    for urlHref in urlHrefARR:
        getImgUrl = goto_content(urlHref)
        del urlHrefARR[0]
        soup = BeautifulSoup(getImgUrl, 'html.parser')
        # 处理图片
        try:
            getImgUrl = soup.find('article', class_='tt-article-content').find_all('img')
            print("洗稿中。。。。")
            time.sleep(random.randint(2, 4))
            html_to_docx(soup, 'index.docx')
            break
        except Exception as e:
            print("链接获取有误正在重新获取")
            continue


# 处理html
def html_to_docx(html_content, output_path, ):
    # 使用BeautifulSoup解析HTML内容
    # soup = BeautifulSoup(html_content, 'html.parser')
    soup = html_content.find('div', class_='main').find('div', class_='article-content')
    # 改写文本内容
    # get_chatgpt(str(soup) + "把里面的文本内容改为英文")

    # 创建一个新的Word文档
    doc = Document()
    imgIndex = 0
    # 遍历HTML中的所有标签，根据需要转换为docx元素
    for element in soup.find_all():
        if element.name == 'p':  # 处理段落
            if element.get_text() != "":
                element_text = get_chatgpt(
                    element.get_text() + " 把空格前面这段文字进行改写,不用返回这一句话'好的，这是改写后的文字：'，类似的回复你也不能返回，你不需要返回你的回复只需要返回结果,直接返回改写后的文字，要求吸引用户眼球，可读性优先，只需要返回改写后的内容，内容中具有代表具体,不用返回这一句话'好的，这是改写后的文字：'，类似的回复你也不能返回，只用返回改写后的文字")
                doc.add_paragraph(element_text)
            elif element.get_text() == "":
                doc.add_paragraph(element.get_text(), )
        elif element.name == 'h1':  # 处理一级标题
            hIndex = 0
            if hIndex == 0:
                if element.get_text() != "":
                    element_text = get_chatgpt(
                        element.get_text() + " 把空格前面这段文字进行改写,不用返回这一句话'好的，这是改写后的文字：'，类似的回复你也不能返回，你不需要返回你的回复只需要返回结果,与原文文字相符程度不能超过10%，包括标点符号,字数要求不能超过30个字，直接返回改写后的文字，要求吸引用户眼球，可读性优先，只需要返回改写后的内容，内容中具有代表具体，不用返回这一句话'好的，这是改写后的文字：'，类似的回复你也不能返回，只用返回改写后的文字")
                    doc.add_heading(element_text, level=1)
                elif element.get_text() == "":
                    doc.add_heading(element.get_text(), level=1)
            if element.get_text() != "":
                element_text = get_chatgpt(
                    element.get_text() + " 把空格前面这段文字进行改写,不用返回这一句话'好的，这是改写后的文字：'，类似的回复你也不能返回，你不需要返回你的回复只需要返回结果,只用返回改写后的文字，与原文文字相符程度不能超过10%，直接返回改写后的文字，要求吸引用户眼球，可读性优先，只需要返回改写后的内容，内容中具有代表具体，")
                doc.add_heading(element_text, level=1)
            elif element.get_text() == "":
                doc.add_heading(element.get_text(), level=1)
            hIndex = hIndex + 1
        # 根据需要添加更多标签的处理逻辑
        # 示例：处理图片（简化示例，实际可能需要更复杂的处理）
        elif element.name == 'img':
            imgType = download_image(element.get('src'), imgIndex)  # 下载图片并保存到本地
            doc.add_picture('./img/' + str(imgIndex) + '.' + imgType, width=Inches(1.25))  # 处理图片
            imgIndex = imgIndex + 1
        # break
    # 保存文档
    doc.save(output_path)
    print('文档生成完成')
    if int(os.path.getsize('index.docx') / 1024 / 1024) >= 15:
        print('生成的文档超过限制，正在重新生成')
        handle_content_img()
    print('开始进行存稿')


# 用户登录获取用户信息接口
def get_user_cookie():
    driver = webdriver.Chrome(options=option, service=service)
    driver.set_window_size(1120, 840)
    try:
        # 打开今日头条主页并手动登录
        driver.get("https://www.toutiao.com/")
        # 等待用户手动登录，直到某个已登录用户才可见的元素出现
        WebDriverWait(driver, 300).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, ".user-icon"))  # 修改为页面中实际存在的已登录元素的选择器
        )
        cookies = driver.get_cookies()
        # 打印当前工作目录
        print("Current working directory: ", os.getcwd())
        # 将 cookie 保存到文件
        with open("cookie/cookies" + str(len(os.listdir('./cookie'))) + ".json", "w") as file:
            json.dump(cookies, file)
            # print(file)
        print("获取成功")
    finally:
        time.sleep(3)
        driver.quit()


# 用户登录接口
def user_login(user, number, cateArr):
    set_sessid(cateArr)  # 获取文章内容
    for i in range(number):
        print("开始第处理" + str(i + 1) + "篇文章")
        handle_content_img()
        write_article(user)
        print("已经保存了" + str(i + 1) + "篇文章")
    time.sleep(10)


def write_article(user):
    driver = webdriver.Chrome(options=option, service=service)
    driver.set_window_size(1120, 840)
    with open("cookie/cookies" + str(user) + ".json", "r") as file:
        cookies = json.load(file)
        # 获取登录后的 cookie
        driver.get("https://www.toutiao.com/")
        driver.headless = True
        # 注入 cookie
        for cookie in cookies:
            driver.add_cookie(cookie)
        time.sleep(5)
        driver.refresh()  # 刷新页面以加载注入的 cookie
        # 这里执行登陆后去写文章的操作
        print("第" + str(user + 1) + "个账号登陆成功")
        time.sleep(5)
    driver.get("https://mp.toutiao.com/profile_v4/graphic/publish")
    print("存稿中耐心等待")
    time.sleep(10)
    try:
        driver.find_element(By.CLASS_NAME, 'drawer-title').find_element(By.TAG_NAME, 'span').click()
    except Exception as e:
        # 处理异常的代码块
        print('没有这个按钮')
    driver.find_element(By.CLASS_NAME, 'doc-import').click()
    # 这里处理写完后的文章
    time.sleep(10)
    wrapper_element = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.CLASS_NAME, "upload-handler"))
    )
    upload_input = wrapper_element.find_element(By.TAG_NAME, "input")
    upload_input.send_keys(os.getcwd() + r"\index.docx")
    time.sleep(45)
    # driver.quit()
    print("草稿储存完成上号查看")
    driver.quit()


# 获取用户收益


# 监听网络信息

def monitor_network():
    if len(os.listdir('./cookie'))==0:
        return
    print("正在获取用户信息")
    all_user_data = []  # 用于存储所有用户的数据
    global user_profit
    options = Options()
    options.add_argument('--headless')  # 不显示浏览器
    caps = {
        "browserName": "chrome",
        'goog:loggingPrefs': {'performance': 'ALL'}  # 开启日志性能监听
    }

    # 将caps添加到options中
    for key, value in caps.items():
        options.set_capability(key, value)
    browser = webdriver.Chrome(options=options)
    for userLen in range(len(os.listdir('./cookie'))):
        with open("cookie/cookies" + str(userLen) + ".json", "r") as file:
            cookies = json.load(file)
            # 获取登录后的 cookie
            browser.get("https://www.toutiao.com/")
            browser.headless = True
            # 注入 cookie
            for cookie in cookies:
                browser.add_cookie(cookie)
            time.sleep(5)
            browser.refresh()  # 刷新页面以加载注入的 cookie
            # 这里执行登陆后去写文章的操作
            print("第" + str(userLen + 1) + "个账号获取成功")
            time.sleep(2)
            browser.get('https://mp.toutiao.com/profile_v4/index')  # 访问该url
            time.sleep(5)

            def filter_type(_type: str):
                types = [
                    'application/javascript', 'application/x-javascript', 'text/css', 'webp', 'image/png', 'image/gif',
                    'image/jpeg', 'image/x-icon', 'application/octet-stream'
                ]
                if _type not in types:
                    return True
                return False

            performance_log = browser.get_log('performance')  # 获取名称为 performance 的日志
            user_info = None
            user_profit = None
            for packet in performance_log:
                message = json.loads(packet.get('message')).get('message')  # 获取message的数据
                if message.get('method') != 'Network.responseReceived':  # 如果method 不是 responseReceived 类型就不往下执行
                    continue
                packet_type = message.get('params').get('response').get('mimeType')  # 获取该请求返回的type
                if not filter_type(_type=packet_type):  # 过滤type
                    continue
                requestId = message.get('params').get('requestId')  # 唯一的请求标识符。相当于该请求的身份证
                url = message.get('params').get('response').get('url')  # 获取 该请求  url
                # print(url)
                try:
                    resp = browser.execute_cdp_cmd('Network.getResponseBody',
                                                   {'requestId': requestId})  # selenium调用 cdp
                    if url == 'https://mp.toutiao.com/mp/fe_api/home/merge_v2?app_id=1231':
                        user_profit = json.loads(resp['body'])['data']['statistic']['data']
                    elif url == 'https://mp.toutiao.com/mp/agw/creator_center/user_info?app_id=1231':
                        user_info = {
                            'user_name': json.loads(resp['body'])['name'],
                            'user_id': json.loads(resp['body'])['user_id'],
                        }
                        # 如果两个链接的数据都已经获取到，就合并它们并存储
                    if user_info and user_profit:
                        user_data = {**user_info, **user_profit}
                        all_user_data.append(user_data)
                        # print(all_user_data)
                        break  # 跳出循环，因为已经获取到所需的数据
                except WebDriverException:  # 忽略异常
                    pass
    return all_user_data


# 开始运行
def start(number, cateArr):
    for userLen in range(len(os.listdir('./cookie'))):
        print("开始第" + str(userLen+1) + "账号")
        user_login(userLen, number, cateArr)


def get_process_id():
    cmd = f"wmic process where name='chromedriver.exe' get processid"
    output = subprocess.check_output(cmd, shell=True, text=True)
    pids = [line.strip() for line in output.split('\n') if line.strip().isdigit()]
    for pid in pids:
        os.system(f'taskkill /F /PID {str(pid)}')
        time.sleep(2)

# 使用方法


if __name__ == "__main__":
    # response = requests.get('https://www.baoww.net/index', )
    # server_cookies = response.cookies
    # server_cookies_dict = requests.utils.dict_from_cookiejar(server_cookies)
    # print(server_cookies_dict['PHPSESSID'])
    print("")
    print(len(os.listdir('./cookie')))
